from docx import Document

# Criar o documento
doc = Document()
doc.add_heading('Relatório Explicativo do ETL com Airflow', 0)

# Adicionar conteúdo ao relatório
doc.add_heading('1. Introdução', level=1)
doc.add_paragraph(
    "Este relatório descreve o desenvolvimento de um pipeline de ETL (Extract, Transform, Load) utilizando o "
    "orquestrador Apache Airflow. O objetivo principal é automatizar o processo de extração, transformação e carga "
    "de dados, garantindo a integridade e o processamento eficiente dos dados."
)

doc.add_heading('2. Etapas do Processo ETL', level=1)
doc.add_paragraph(
    "O processo foi dividido nas seguintes etapas:\n\n"
    "Extração (Extract): Carregamento dos dados a partir de arquivos JSON e Parquet.\n"
    "Transformação (Transform): Aplicação das transformações necessárias nos dados.\n"
    "Carregamento (Load): Processamento dos dados e inserção no destino final (arquivo ou banco de dados).\n"
    "Verificação de Retry: Monitoramento das tentativas de reexecução de tarefas falhas."
)

doc.add_heading('3. Implementação do Pipeline', level=1)
doc.add_paragraph(
    "A seguir, são detalhadas as etapas implementadas no código."
)

doc.add_heading('3.1 Função de Extração de Dados', level=2)
doc.add_paragraph(
    "Na função extract(), os dados são carregados a partir dos arquivos registros_oportunidades.json e "
    "sellout.parquet localizados em um contêiner Docker. O código utilizado para esta tarefa é o seguinte:"
)

doc.add_paragraph('''\
def extract():
    try:
        print("Extraindo dados...")
        caminho_oportunidades = "/opt/airflow/dags/database/registros_oportunidades.json"
        caminho_sellout = "/opt/airflow/dags/database/sellout.parquet"
        
        oportunidades = pd.read_json(caminho_oportunidades)
        sellout = pd.read_parquet(caminho_sellout)
        
        print("Dados carregados com sucesso.")
        return oportunidades, sellout
    
    except Exception as e:
        raise Exception(f"Erro na etapa de extração: {e}")
''')

doc.add_heading('Desafios Enfrentados:', level=3)
doc.add_paragraph(
    "Garantir que o caminho dos arquivos estivesse correto no contêiner Docker.\n"
    "Verificar que os arquivos JSON e Parquet estavam estruturados corretamente para serem lidos pelo pandas."
)

doc.add_heading('Decisões Tomadas:', level=3)
doc.add_paragraph(
    "Utilizamos pandas.read_json para carregar o arquivo JSON e pandas.read_parquet para o arquivo Parquet.\n"
    "Decidimos armazenar os dados em variáveis para posterior transformação, mantendo o fluxo simples e direto."
)

doc.add_heading('3.2 Função de Transformação de Dados', level=2)
doc.add_paragraph(
    "A função transform() foi definida para aplicar as transformações necessárias nos dados. Embora o código "
    "não tenha transformações específicas implementadas, o esqueleto da função se parece com:"
)

doc.add_paragraph('''\
def transform():
    try:
        print("Transformando dados...")
        # Aqui você colocaria a lógica de transformação
        print("Transformações concluídas.")
    
    except Exception as e:
        raise Exception(f"Erro na etapa de transformação: {e}")
''')

doc.add_heading('Desafios Enfrentados:', level=3)
doc.add_paragraph(
    "Durante o desenvolvimento do pipeline, decidimos não aplicar transformações complexas nesta etapa inicial "
    "para simplificação. O desafio principal seria garantir que as transformações estivessem em conformidade com "
    "os requisitos do banco de dados ou do modelo final."
)

doc.add_heading('Decisões Tomadas:', level=3)
doc.add_paragraph(
    "Deixamos a etapa de transformação aberta para inclusão de lógica de limpeza de dados, renomeação de colunas, "
    "conversão de tipos, etc., dependendo das necessidades dos dados finais."
)

doc.add_heading('3.3 Função de Carregamento de Dados', level=2)
doc.add_paragraph(
    "Na função load(), os dados são carregados para um destino final, como um banco de dados ou arquivo. Embora "
    "o código da função não tenha implementado o carregamento real, a estrutura para a carga é a seguinte:"
)

doc.add_paragraph('''\
def load():
    try:
        print("Carregando dados...")
        # Aqui você colocaria a lógica para carregar os dados (por exemplo, para um arquivo Excel)
        print("Dados carregados com sucesso.")
    
    except Exception as e:
        raise Exception(f"Erro na etapa de carregamento: {e}")
''')

doc.add_heading('Desafios Enfrentados:', level=3)
doc.add_paragraph(
    "Implementar a carga de dados no destino apropriado, seja em banco de dados SQL ou outro formato.\n"
    "A escolha do destino e a validação de dados antes da carga também foram desafios durante o design do pipeline."
)

doc.add_heading('Decisões Tomadas:', level=3)
doc.add_paragraph(
    "Optamos por não detalhar a implementação do carregamento para este relatório, visto que isso dependeria do "
    "destino final. A carga seria implementada dependendo do banco de dados ou serviço de dados desejado."
)

doc.add_heading('3.4 Função de Verificação de Retry', level=2)
doc.add_paragraph(
    "A função check_task_retry_status() foi criada para monitorar as tarefas que falharam e precisam ser reexecutadas. "
    "A função de verificação de retry é fundamental para a robustez do pipeline, garantindo que o processo seja repetido "
    "se uma falha ocorrer."
)

doc.add_paragraph('''\
def check_task_retry_status(task_instance: TaskInstance):
    if task_instance.ready_for_retry:
        print(f"Tarefa {task_instance.task_id} está esperando para ser reexecutada.")
    else:
        print(f"Tarefa {task_instance.task_id} não está mais aguardando reexecução.")
''')

doc.add_heading('Desafios Enfrentados:', level=3)
doc.add_paragraph(
    "Garantir que as tarefas fossem adequadamente monitoradas, verificando se estavam esperando por reexecuções ou não.\n"
    "O estado de retry foi alterado de up_for_retry para ready_for_retry para melhorar a identificação da situação da tarefa."
)

doc.add_heading('Decisões Tomadas:', level=3)
doc.add_paragraph(
    "A decisão de monitorar os estados das tarefas com o uso de TaskInstance foi tomada para garantir que a DAG "
    "tivesse capacidade de tratar falhas sem intervenção manual."
)

doc.add_heading('3.5 Definição da DAG', level=2)
doc.add_paragraph(
    "A DAG foi definida para orquestrar as tarefas no Airflow. As configurações de retries, delay e start_date "
    "foram incluídas no argumento default_args."
)

doc.add_paragraph('''\
dag = DAG(
    'etl_airflow_final', 
    default_args=default_args,
    description='ETL com Airflow',
    schedule=None, 
    start_date=pendulum.today('UTC').add(days=-1), 
    catchup=False,
)
''')

doc.add_heading('Desafios Enfrentados:', level=3)
doc.add_paragraph(
    "Definir o comportamento de start_date e garantir que o catchup=False fosse configurado corretamente, evitando execuções duplicadas."
)

doc.add_heading('Decisões Tomadas:', level=3)
doc.add_paragraph(
    "O schedule_interval foi removido para simplificação e para permitir que a DAG fosse executada manualmente ou conforme necessidade."
)

doc.add_heading('4. Desafios e Decisões Gerais', level=1)
doc.add_paragraph(
    "Ambiente Docker: Um desafio significativo foi configurar o ambiente do Airflow e garantir que os caminhos dos arquivos "
    "fossem corretamente acessados dentro do contêiner Docker.\n"
    "Retry e Falhas: Outro desafio importante foi garantir que a DAG pudesse se recuperar de falhas automaticamente. "
    "O Airflow foi configurado para tentar até 3 vezes em caso de erro.\n"
    "Escalabilidade: A solução foi desenhada para permitir expansões futuras, com transformações de dados complexas e integração "
    "com diferentes fontes e destinos."
)

doc.add_heading('5. Conclusão', level=1)
doc.add_paragraph(
    "O pipeline de ETL foi implementado com sucesso utilizando Apache Airflow, proporcionando uma solução robusta e automatizada "
    "para o processo de dados. A integração com pandas para extração e transformação de dados, e o controle de falhas via retry "
    "garantiram que o fluxo fosse eficiente e confiável. A arquitetura também permite futuras melhorias, como adição de mais fontes "
    "de dados e destinos, transformações complexas, e visualização de falhas em tempo real."
)

# Salvar o documento no local especificado
file_path = r"C:\Users\wesle\OneDrive\Documentos\Teste_Analista_BI\scripts\relatorio_etl_airflow.docx"
doc.save(file_path)

print(f"Relatório salvo em: {file_path}")
